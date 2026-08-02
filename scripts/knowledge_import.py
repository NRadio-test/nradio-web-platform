#!/usr/bin/env python3
"""Parse an uploaded document and turn it into reviewable NRadio knowledge entries."""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
import re
import shutil
import sys
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any


SUPPORTED_SUFFIXES = {
    ".pdf", ".txt", ".md", ".markdown", ".rst", ".adoc", ".docx", ".xlsx", ".xls", ".epub"
}
CONFIDENCE_VALUES = {"high", "medium_high", "medium", "low_medium"}
DECISION_VALUES = {"accept", "needs_review", "reject"}
STRUCTURE_PROMPT = """你是 NRadio 鲲鹏无限知识库的内容结构化编辑。上传者已经通过成员身份验证，并确认有权使用资料；你的职责是忠实提取、整理和拆分知识，不是判断资料是否值得收录。
只输出 JSON 对象，不得输出 Markdown。结构必须为：
{"decision":"accept|needs_review|reject","review_notes":["..."],"entries":[{"title":"...","text":"...","tags":["..."],"confidence":"high|medium_high|medium|low_medium"}]}
规则：
1. 默认 decision=accept，并尽可能生成至少一条 entries；只有内容存在明显矛盾、含义不清或需要成员核对时使用 needs_review，但仍须保留可以提取的条目。
2. 上传内容均由已验证成员主动选择并确认可以写入知识库。不得进行隐私、安全性、内部资料或主题相关性筛选，不得删除、脱敏或拒绝其中的信息。
3. 联系人姓名、手机号、QQ、微信、邮箱、群号、直播间、内部通知、售后流程和业务规则等都应按照原文保留并转换为可检索知识。
4. 仅当文本为空、乱码、无法理解或完全没有任何可陈述信息时，才可 decision=reject 且 entries=[]。不要因为内容与产品参数无关而拒绝。
5. 每个条目只表达一个便于检索和回答的主题；text 使用完整中文，并保留日期、适用对象、型号后缀、条件、例外及联系方式。
6. 价格、库存、活动、覆盖、测速、账号统计等时效性内容可以收录，但须在正文中保留资料给出的日期或时间边界。
7. 不要补写资料中不存在的事实；只做总结、改写、拆分、去除无意义重复和补充检索标签。"""


def safe_name(value: str) -> str:
    name = Path(value.replace("\\", "/")).name
    cleaned = re.sub(r"[^\w. -]+", "_", name, flags=re.UNICODE).lstrip(".")
    return cleaned[:160] or "document"


def resolve_source_url(reference: str, repository: str, source_path: Path) -> str:
    value = reference.strip()
    if re.match(r"^https?://[^\s]+$", value, flags=re.IGNORECASE):
        return value
    return f"https://github.com/{repository}/blob/main/{source_path.as_posix()}"


def read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "big5"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def extract_pdf(path: Path) -> str:
    import fitz

    document = fitz.open(path)
    return "\n\n".join(page.get_text("text") for page in document)


def extract_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n\n".join(blocks)


def extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    output: list[str] = []
    for sheet in workbook.worksheets:
        output.append(f"# 工作表：{sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() if value is not None else "" for value in row]
            if any(values):
                output.append(" | ".join(values))
    return "\n".join(output)


def extract_xls(path: Path) -> str:
    import xlrd

    workbook = xlrd.open_workbook(path)
    output: list[str] = []
    for sheet in workbook.sheets():
        output.append(f"# 工作表：{sheet.name}")
        for row_index in range(sheet.nrows):
            values = [str(sheet.cell_value(row_index, column)).strip() for column in range(sheet.ncols)]
            if any(values):
                output.append(" | ".join(values))
    return "\n".join(output)


def extract_epub(path: Path) -> str:
    from bs4 import BeautifulSoup
    from ebooklib import ITEM_DOCUMENT, epub

    book = epub.read_epub(str(path))
    output: list[str] = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        output.append(BeautifulSoup(item.get_content(), "html.parser").get_text("\n"))
    return "\n\n".join(output)


def extract_text(path: Path, original_name: str) -> str:
    suffix = Path(original_name).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"暂不支持的文件格式：{suffix or '无扩展名'}")
    if suffix == ".pdf":
        text = extract_pdf(path)
    elif suffix == ".docx":
        text = extract_docx(path)
    elif suffix == ".xlsx":
        text = extract_xlsx(path)
    elif suffix == ".xls":
        text = extract_xls(path)
    elif suffix == ".epub":
        text = extract_epub(path)
    else:
        text = read_text_file(path)
    text = re.sub(r"\x00", "", text).strip()
    if len(text) < 40:
        raise ValueError("未能提取足够文本；扫描版或纯图片 PDF 需要先进行 OCR。")
    return text


def chunk_text(text: str, chunk_size: int = 24_000, max_chunks: int = 24) -> list[str]:
    chunks: list[str] = []
    cursor = 0
    while cursor < len(text) and len(chunks) < max_chunks:
        end = min(len(text), cursor + chunk_size)
        if end < len(text):
            split = max(text.rfind("\n", cursor, end), text.rfind("。", cursor, end))
            if split > cursor + chunk_size // 2:
                end = split + 1
        chunks.append(text[cursor:end])
        cursor = end
    return chunks


def request_review(document_name: str, text: str, part: int, total: int, notes: str) -> dict[str, Any]:
    api_key = os.environ.get("KNOWLEDGE_REVIEW_API_KEY", "").strip()
    api_base = os.environ.get("KNOWLEDGE_REVIEW_API_BASE", "https://api.openai.com/v1").rstrip("/")
    model = os.environ.get("KNOWLEDGE_REVIEW_MODEL", "").strip()
    if not api_key or not model:
        raise RuntimeError("未配置 KNOWLEDGE_REVIEW_API_KEY 或 KNOWLEDGE_REVIEW_MODEL。")

    user_prompt = (
        f"文件：{document_name}\n分段：{part}/{total}\n上传者说明：{notes or '无'}\n\n"
        f"待审核文本：\n{text}"
    )
    payload = {
        "model": model,
        "temperature": 0.1,
        "response_format": {"type": "json_object"},
        "messages": [
            {"role": "system", "content": STRUCTURE_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
    }
    request = urllib.request.Request(
        f"{api_base}/chat/completions",
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "Accept": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=180) as response:
            body = json.load(response)
    except urllib.error.HTTPError as error:
        detail = error.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"结构化模型返回 {error.code}：{detail[:500]}") from error
    content = body["choices"][0]["message"]["content"].strip()
    content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content)
    return json.loads(content)


def validate_review(payload: dict[str, Any]) -> dict[str, Any]:
    decision = payload.get("decision")
    if decision not in DECISION_VALUES:
        raise ValueError("审核结果缺少有效 decision。")
    notes = [str(item).strip()[:500] for item in payload.get("review_notes", []) if str(item).strip()]
    entries: list[dict[str, Any]] = []
    for raw in payload.get("entries", []):
        title = str(raw.get("title", "")).strip()[:160]
        text = str(raw.get("text", "")).strip()[:4000]
        tags = list(dict.fromkeys(str(item).strip()[:40] for item in raw.get("tags", []) if str(item).strip()))[:12]
        confidence = str(raw.get("confidence", "medium")).strip()
        if title and len(text) >= 20:
            entries.append({
                "title": title,
                "text": text,
                "tags": tags or ["待分类"],
                "confidence": confidence if confidence in CONFIDENCE_VALUES else "medium",
            })
    if entries and decision == "reject":
        decision = "needs_review"
    return {"decision": decision, "review_notes": notes, "entries": entries}


def write_outputs(args: argparse.Namespace, extracted: str, reviews: list[dict[str, Any]]) -> dict[str, Any]:
    today = dt.date.today().isoformat()
    digest = hashlib.sha256(Path(args.input).read_bytes()).hexdigest()
    short_digest = digest[:10]
    output_root = Path(args.output_root).resolve()
    source_dir = output_root / "knowledge-base" / "sources" / "uploads" / today[:7]
    document_dir = output_root / "knowledge-base" / "documents" / "uploads"
    review_dir = output_root / "knowledge-base" / "reviews"
    source_dir.mkdir(parents=True, exist_ok=True)
    document_dir.mkdir(parents=True, exist_ok=True)
    review_dir.mkdir(parents=True, exist_ok=True)

    filename = safe_name(args.file_name)
    uploaded_by = args.uploaded_by or "unknown"
    source_path = source_dir / f"{args.job_id}-{filename}"
    shutil.copyfile(args.input, source_path)
    relative_source_path = source_path.relative_to(output_root)
    source_url = resolve_source_url(
        args.source_url,
        args.github_repository,
        relative_source_path,
    )

    review_notes = list(dict.fromkeys(note for item in reviews for note in item["review_notes"]))
    raw_entries = [entry for item in reviews for entry in item["entries"]]
    unique: dict[str, dict[str, Any]] = {}
    for entry in raw_entries:
        key = re.sub(r"\s+", "", entry["title"]).lower()
        unique.setdefault(key, entry)

    decision = (
        "reject"
        if not unique
        else "needs_review"
        if any(item["decision"] == "needs_review" for item in reviews)
        else "accept"
    )
    final_entries: list[dict[str, Any]] = []
    for index, entry in enumerate(unique.values(), start=1):
        final_entries.append({
            "id": f"upload-{today.replace('-', '')}-{short_digest}-{index:02d}",
            "title": entry["title"],
            "text": entry["text"],
            "source_url": source_url,
            "source_type": "user_upload",
            "uploaded_by": uploaded_by,
            "verified_at": today,
            "confidence": entry["confidence"],
            "tags": entry["tags"],
        })

    document_path = document_dir / f"{args.job_id}.md"
    lines = [
        f"# {args.title or filename}", "",
        f"- 导入任务：`{args.job_id}`",
        f"- 原始文件：`{source_path.relative_to(output_root).as_posix()}`",
        f"- SHA-256：`{digest}`",
        f"- 审核结论：`{decision}`",
        f"- 原始来源：{args.source_url.strip() or '未提供'}",
        f"- 上传者：{uploaded_by}",
        "",
        "## 审核备注", "",
        *([f"- {note}" for note in review_notes] or ["- 无额外备注。"]),
        "", "## 结构化知识草稿", "",
    ]
    for entry in final_entries:
        lines.extend([
            f"### {entry['title']}", "", entry["text"], "",
            f"标签：{', '.join(entry['tags'])}", "",
        ])
    if not final_entries:
        lines.append("本次审核未生成可进入正式知识库的条目。")
    document_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")

    review_payload = {
        "job_id": args.job_id,
        "file_name": filename,
        "uploaded_by": uploaded_by,
        "sha256": digest,
        "decision": decision,
        "review_notes": review_notes,
        "entry_count": len(final_entries),
        "document_path": document_path.relative_to(output_root).as_posix(),
        "source_path": source_path.relative_to(output_root).as_posix(),
        "extracted_characters": len(extracted),
    }
    (review_dir / f"{args.job_id}.json").write_text(
        json.dumps(review_payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )

    if final_entries:
        jsonl_path = output_root / "knowledge-base" / "import" / "knowledge.jsonl"
        jsonl_path.parent.mkdir(parents=True, exist_ok=True)
        with jsonl_path.open("a", encoding="utf-8") as handle:
            for entry in final_entries:
                handle.write(json.dumps(entry, ensure_ascii=False, separators=(",", ":")) + "\n")
    return review_payload


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--file-name", required=True)
    parser.add_argument("--job-id", required=True)
    parser.add_argument("--output-root", required=True)
    parser.add_argument("--github-repository", default="NRadio-test/nradio-web-platform")
    parser.add_argument("--title", default="")
    parser.add_argument("--source-url", default="")
    parser.add_argument("--notes", default="")
    parser.add_argument("--uploaded-by", default="")
    parser.add_argument("--metadata", default="")
    args = parser.parse_args()
    if args.metadata:
        payload = json.loads(Path(args.metadata).read_text(encoding="utf-8"))
        metadata = payload.get("job", payload)
        args.file_name = str(metadata.get("filename") or args.file_name)
        args.title = str(metadata.get("title") or args.title)
        args.source_url = str(metadata.get("source_url") or args.source_url)
        args.notes = str(metadata.get("notes") or args.notes)
        args.uploaded_by = str(metadata.get("uploaded_by") or args.uploaded_by)
    return args


def main() -> int:
    args = parse_args()
    extracted = extract_text(Path(args.input), args.file_name)
    chunks = chunk_text(extracted)
    reviews = [
        validate_review(request_review(args.file_name, chunk, index, len(chunks), args.notes))
        for index, chunk in enumerate(chunks, start=1)
    ]
    result = write_outputs(args, extracted, reviews)
    print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as error:
        print(f"knowledge import failed: {error}", file=sys.stderr)
        raise
